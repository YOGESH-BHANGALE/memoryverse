"""
File parsers — extract raw text from PDF, DOCX, and TXT files.
Each parser returns a normalised `RawDocument`.
"""

from __future__ import annotations

import io
from pathlib import Path

from app.models.schemas import FileType, RawDocument
from app.utils.logger import logger


# ── PDF Parser ──────────────────────────────────────────────────────────

class PDFParser:
    """Extract text from PDF files using PyPDF2 with pdfplumber fallback."""

    @staticmethod
    def parse(file_bytes: bytes, filename: str) -> RawDocument:
        # PyPDF2 raises (not just returns empty) on encrypted or malformed
        # PDFs, which previously escaped as an HTTP 500 instead of using the
        # pdfplumber fallback this class documents.
        try:
            text, page_count = PDFParser._try_pypdf2(file_bytes)
        except Exception as exc:
            logger.warning(f"PyPDF2 failed to parse {filename} ({exc}); falling back to pdfplumber")
            text, page_count = "", 0

        if not text.strip():
            logger.info("PyPDF2 returned empty text, falling back to pdfplumber")
            text, page_count = PDFParser._try_pdfplumber(file_bytes)
        return RawDocument(
            text=text,
            filename=filename,
            file_type=FileType.PDF,
            page_count=page_count,
        )

    @staticmethod
    def _try_pypdf2(data: bytes) -> tuple[str, int]:
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(data))
        pages: list[str] = []
        for page in reader.pages:
            extracted = page.extract_text() or ""
            pages.append(extracted)
        return "\n\n".join(pages), len(reader.pages)

    @staticmethod
    def _try_pdfplumber(data: bytes) -> tuple[str, int]:
        import pdfplumber

        pages: list[str] = []
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    pages.append(text)
                return "\n\n".join(pages), len(pdf.pages)
        except Exception as exc:
            # Encrypted or corrupt PDFs make pdfplumber.open raise. PyPDF2 is
            # already exhausted by the time we get here, so there is nothing
            # left to try — return empty text and let the caller reject the
            # upload with a clean 400 rather than the exception escaping as a
            # 500.
            logger.warning(
                f"pdfplumber also failed to parse the PDF ({exc}); giving up with empty text"
            )
            return "", 0


# ── Text Parser ─────────────────────────────────────────────────────────

class TextParser:
    """Read plain text files."""

    @staticmethod
    def parse(file_bytes: bytes, filename: str) -> RawDocument:
        text = file_bytes.decode("utf-8", errors="replace")
        return RawDocument(
            text=text,
            filename=filename,
            file_type=FileType.TXT,
            page_count=1,
        )


# ── DOCX Parser ─────────────────────────────────────────────────────────

class DocxParser:
    """Extract text from DOCX files using python-docx."""

    @staticmethod
    def parse(file_bytes: bytes, filename: str) -> RawDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        return RawDocument(
            text=text,
            filename=filename,
            file_type=FileType.DOCX,
            page_count=max(1, len(paragraphs) // 25),  # rough estimate
        )


# ── Router ───────────────────────────────────────────────────────────────

_PARSERS = {
    FileType.PDF: PDFParser.parse,
    FileType.TXT: TextParser.parse,
    FileType.DOCX: DocxParser.parse,
}


def parse_file(file_bytes: bytes, filename: str, file_type: FileType) -> RawDocument:
    """Route to the correct parser based on file type."""
    parser_fn = _PARSERS.get(file_type)
    if not parser_fn:
        raise ValueError(f"No parser registered for file type: {file_type}")
    logger.info(f"Parsing {filename} as {file_type.value}")
    return parser_fn(file_bytes, filename)
